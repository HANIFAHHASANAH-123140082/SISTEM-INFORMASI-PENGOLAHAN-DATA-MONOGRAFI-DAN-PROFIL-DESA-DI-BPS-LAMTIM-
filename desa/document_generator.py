import os
from datetime import date
from PIL import Image, ImageDraw, ImageFont, ImageFilter

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import get_object_or_404

from .models import Desa, Bab, JawabanKuesioner
from .utils import bersihkan_teks_tempel

# ==================== PATH ASSET (logo, cover, font) ====================

ASSETS_DIR = os.path.join(settings.BASE_DIR, 'desa', 'assets')
LOGO_PATH = os.path.join(ASSETS_DIR, 'logo_lampung_timur.png')       # <- logo dipakai di kop surat CV & cover Monografi/Profil
COVERS_DIR = os.path.join(ASSETS_DIR, 'covers')
GENERATED_DIR = os.path.join(settings.MEDIA_ROOT, 'generated_covers')
os.makedirs(GENERATED_DIR, exist_ok=True)

FONT_BOLD = "C:/Windows/Fonts/timesbd.ttf"
FONT_REGULAR = "C:/Windows/Fonts/times.ttf"

DPI = 150
PAGE_W_CM, PAGE_H_CM = 21.0, 29.7
PAGE_W_PX = int(PAGE_W_CM / 2.54 * DPI)
PAGE_H_PX = int(PAGE_H_CM / 2.54 * DPI)

BULAN_INDO = {
    1: "Januari", 2: "Februari", 3: "Maret", 4: "April", 5: "Mei", 6: "Juni",
    7: "Juli", 8: "Agustus", 9: "September", 10: "Oktober", 11: "November", 12: "Desember"
}


def tanggal_indonesia():
    today = date.today()
    return f"{BULAN_INDO[today.month]} {today.year}"


def _font(path, size):
    try:
        return ImageFont.truetype(path, size)
    except Exception:
        return ImageFont.load_default()


def _crop_to_cover(img, target_w, target_h):
    src_ratio = img.width / img.height
    target_ratio = target_w / target_h
    if src_ratio > target_ratio:
        new_height = img.height
        new_width = int(new_height * target_ratio)
        left = (img.width - new_width) // 2
        img = img.crop((left, 0, left + new_width, new_height))
    else:
        new_width = img.width
        new_height = int(new_width / target_ratio)
        top = (img.height - new_height) // 2
        img = img.crop((0, top, new_width, top + new_height))
    return img.resize((target_w, target_h), Image.LANCZOS)


def get_cover_image_path(desa):
    if not os.path.isdir(COVERS_DIR):
        return None
    daftar = sorted([f for f in os.listdir(COVERS_DIR) if f.lower().endswith(('.jpg', '.jpeg', '.png'))])
    if not daftar:
        return None
    return os.path.join(COVERS_DIR, daftar[desa.id % len(daftar)])


# ==================== GENERATE GAMBAR COVER DEPAN (Profil Desa) ====================

def generate_cover_depan(desa):
    src_path = get_cover_image_path(desa)
    canvas = Image.new('RGB', (PAGE_W_PX, PAGE_H_PX), '#FFFFFF')

    if src_path:
        img = Image.open(src_path).convert('RGB')
        img = _crop_to_cover(img, PAGE_W_PX, PAGE_H_PX)
        canvas.paste(img, (0, 0))

    draw = ImageDraw.Draw(canvas, 'RGBA')
    pad = int(0.06 * PAGE_W_PX)  # margin kiri-kanan konsisten

    # ---- LANGKAH 1: gradasi gelap tipis & blur di area atas ----
    band_top_h = int(0.15 * PAGE_H_PX)
    overlay_atas = Image.new('RGBA', (PAGE_W_PX, band_top_h), (0, 0, 0, 0))
    draw_overlay = ImageDraw.Draw(overlay_atas)
    for y in range(band_top_h):
        alpha = int(90 * (1 - y / band_top_h))
        draw_overlay.line([(0, y), (PAGE_W_PX, y)], fill=(0, 0, 0, alpha))
    overlay_atas = overlay_atas.filter(ImageFilter.GaussianBlur(radius=int(0.01 * PAGE_H_PX)))
    canvas.paste(overlay_atas, (0, 0), overlay_atas)
    draw = ImageDraw.Draw(canvas, 'RGBA')

    # ---- LANGKAH 2: dekorasi lengkung emas pojok kanan atas ----
    swirl_box = int(0.20 * PAGE_W_PX)
    for i, lebar in enumerate([3, 2, 1]):
        inset = i * int(0.015 * PAGE_W_PX)
        draw.arc(
            [PAGE_W_PX - swirl_box + inset, -swirl_box // 2 + inset,
             PAGE_W_PX + swirl_box // 3, swirl_box // 2 + inset],
            start=135, end=230, fill='#D9A44C', width=lebar
        )

    # ---- LANGKAH 3: logo + teks instansi, dihitung presisi & rapat ----
    logo_h = int(0.11 * PAGE_H_PX)   # <- UBAH DI SINI kalau logo cover depan mau lebih besar/kecil
    logo_y = int(0.028 * PAGE_H_PX)   # jarak dari tepi atas, tidak berlebihan
    teks_x = pad

    if os.path.isfile(LOGO_PATH):
        logo = Image.open(LOGO_PATH).convert('RGBA')
        logo_w = int(logo.width * (logo_h / logo.height))
        logo_resized = logo.resize((logo_w, logo_h), Image.LANCZOS)
        canvas.paste(logo_resized, (pad, logo_y), logo_resized)
        teks_x = pad + logo_w + int(0.008 * PAGE_W_PX)  # jarak logo ke teks dirapatkan

    lebar_maks_teks = PAGE_W_PX - teks_x - pad

    def _cari_ukuran_pas(teks, ukuran_awal, font_path):
        ukuran = ukuran_awal
        while ukuran > int(0.008 * PAGE_H_PX):
            font = _font(font_path, ukuran)
            lebar = draw.textbbox((0, 0), teks, font=font)[2]
            if lebar <= lebar_maks_teks:
                return font
            ukuran -= 1
        return _font(font_path, int(0.008 * PAGE_H_PX))

    f_pemda = _cari_ukuran_pas("PEMERINTAH KABUPATEN LAMPUNG TIMUR", int(0.016 * PAGE_H_PX), FONT_BOLD)
    f_desa = _cari_ukuran_pas(f"Pemerintah Desa {desa.nama_desa}", int(0.013 * PAGE_H_PX), FONT_REGULAR)

    tinggi_pemda = f_pemda.getbbox("PEMERINTAH")[3]
    tinggi_desa = f_desa.getbbox("Pemerintah")[3]
    jarak_antar_baris = int(0.008 * PAGE_H_PX)
    tinggi_teks_total = tinggi_pemda + jarak_antar_baris + tinggi_desa

    # Teks disejajarkan vertikal tepat di tengah tinggi logo
    teks_y = logo_y + (logo_h - tinggi_teks_total) // 2

    draw.text((teks_x, teks_y), "PEMERINTAH KABUPATEN LAMPUNG TIMUR", font=f_pemda, fill='white')
    draw.text((teks_x, teks_y + tinggi_pemda + jarak_antar_baris), f"Pemerintah Desa {desa.nama_desa}", font=f_desa, fill='#EDEDED')

    # ---- LANGKAH 4: panel judul di tengah halaman ----
    f_label = _font(FONT_REGULAR, int(0.024 * PAGE_H_PX))
    f_nama = _font(FONT_BOLD, int(0.052 * PAGE_H_PX))
    f_tahun = _font(FONT_BOLD, int(0.022 * PAGE_H_PX))

    label = "P R O F I L   D E S A"
    nama_desa_teks = desa.nama_desa.upper()
    tahun_teks = f"T A H U N   {desa.tahun_profil}"

    bbox_label = draw.textbbox((0, 0), label, font=f_label)
    bbox_nama = draw.textbbox((0, 0), nama_desa_teks, font=f_nama)
    bbox_tahun = draw.textbbox((0, 0), tahun_teks, font=f_tahun)

    label_w, label_h = bbox_label[2] - bbox_label[0], bbox_label[3] - bbox_label[1]
    nama_w, nama_h = bbox_nama[2] - bbox_nama[0], bbox_nama[3] - bbox_nama[1]
    tahun_w, tahun_h = bbox_tahun[2] - bbox_tahun[0], bbox_tahun[3] - bbox_tahun[1]

    ornamen_h = int(0.03 * PAGE_H_PX)
    jarak1 = int(0.015 * PAGE_H_PX)
    jarak2 = int(0.02 * PAGE_H_PX)
    garis_h = int(0.02 * PAGE_H_PX)
    jarak3 = int(0.02 * PAGE_H_PX)

    tinggi_konten = ornamen_h + jarak1 + label_h + jarak2 + nama_h + jarak2 + garis_h + jarak3 + tahun_h
    panel_pad_v = int(0.035 * PAGE_H_PX)
    panel_pad_h = int(0.09 * PAGE_W_PX)

    panel_w = min(max(nama_w, label_w) + panel_pad_h * 2, int(0.85 * PAGE_W_PX))
    panel_h = tinggi_konten + panel_pad_v * 2

    panel_x0 = (PAGE_W_PX - panel_w) // 2
    panel_y0 = (PAGE_H_PX - panel_h) // 2
    panel_x1 = panel_x0 + panel_w
    panel_y1 = panel_y0 + panel_h

    panel_layer = Image.new('RGBA', (PAGE_W_PX, PAGE_H_PX), (0, 0, 0, 0))
    panel_draw = ImageDraw.Draw(panel_layer)
    radius = int(0.02 * PAGE_H_PX)
    panel_draw.rounded_rectangle(
        [panel_x0, panel_y0, panel_x1, panel_y1],
        radius=radius, fill=(10, 35, 22, 190), outline='#D9A44C', width=2
    )
    canvas.paste(panel_layer, (0, 0), panel_layer)
    draw = ImageDraw.Draw(canvas, 'RGBA')

    cx = PAGE_W_PX // 2
    diamond_r = int(0.006 * PAGE_H_PX)

    cy = panel_y0 + panel_pad_v + ornamen_h // 2
    draw.line([(cx - int(0.05 * PAGE_W_PX), cy), (cx - diamond_r * 2, cy)], fill='#D9A44C', width=2)
    draw.line([(cx + diamond_r * 2, cy), (cx + int(0.05 * PAGE_W_PX), cy)], fill='#D9A44C', width=2)
    draw.polygon(
        [(cx, cy - diamond_r), (cx + diamond_r, cy), (cx, cy + diamond_r), (cx - diamond_r, cy)],
        fill='#D9A44C'
    )

    y_cursor = panel_y0 + panel_pad_v + ornamen_h + jarak1
    draw.text(((PAGE_W_PX - label_w) / 2, y_cursor), label, font=f_label, fill='#EDEDED')

    y_cursor += label_h + jarak2
    draw.text(((PAGE_W_PX - nama_w) / 2, y_cursor), nama_desa_teks, font=f_nama, fill='white')

    y_cursor += nama_h + jarak2
    garis_w_deco = int(0.10 * PAGE_W_PX)
    garis_y = y_cursor + garis_h // 2
    draw.line([(cx - garis_w_deco, garis_y), (cx - diamond_r * 2, garis_y)], fill='#D9A44C', width=2)
    draw.line([(cx + diamond_r * 2, garis_y), (cx + garis_w_deco, garis_y)], fill='#D9A44C', width=2)
    draw.polygon(
        [(cx, garis_y - diamond_r), (cx + diamond_r, garis_y), (cx, garis_y + diamond_r), (cx - diamond_r, garis_y)],
        fill='#D9A44C'
    )

    y_cursor += garis_h + jarak3
    draw.text(((PAGE_W_PX - tahun_w) / 2, y_cursor), tahun_teks, font=f_tahun, fill='#D9A44C')

    out_path = os.path.join(GENERATED_DIR, f"cover_depan_{desa.id}.png")
    canvas.save(out_path)
    return out_path


def _gambar_teks_dengan_halo(draw, pos, teks, font, fill, warna_halo='#FFFFFF', kekuatan=None):
    """
    Gambar teks dengan 'halo' tipis di sekelilingnya (beberapa salinan teks digeser
    sedikit ke segala arah dengan warna terang) supaya teks gelap tetap kebaca
    walau ditumpuk di atas foto yang terang, tanpa perlu bikin panel solid.
    """
    x, y = pos
    if kekuatan is None:
        kekuatan = max(2, font.size // 20)
    arah = [(-kekuatan, 0), (kekuatan, 0), (0, -kekuatan), (0, kekuatan),
            (-kekuatan, -kekuatan), (kekuatan, -kekuatan), (-kekuatan, kekuatan), (kekuatan, kekuatan)]
    for dx, dy in arah:
        draw.text((x + dx, y + dy), teks, font=font, fill=warna_halo)
    draw.text((x, y), teks, font=font, fill=fill)


# ==================== GENERATE GAMBAR COVER BELAKANG (Profil Desa) ====================

def generate_cover_belakang(desa):
    src_path = get_cover_image_path(desa)
    canvas = Image.new('RGB', (PAGE_W_PX, PAGE_H_PX), '#FFFFFF')

    if src_path:
        img = Image.open(src_path).convert('RGB')
        img = _crop_to_cover(img, PAGE_W_PX, PAGE_H_PX)
        canvas.paste(img, (0, 0))

    draw = ImageDraw.Draw(canvas, 'RGBA')
    pad = int(0.06 * PAGE_W_PX)

    # ---- gradasi gelap di bagian bawah (menaungi teks "Terima Kasih" + logo) ----
    band_bawah_h = int(0.36 * PAGE_H_PX)
    band_bawah_y0 = PAGE_H_PX - band_bawah_h
    overlay_bawah = Image.new('RGBA', (PAGE_W_PX, band_bawah_h), (0, 0, 0, 0))
    draw_overlay = ImageDraw.Draw(overlay_bawah)
    for y in range(band_bawah_h):
        alpha = int(140 * (y / band_bawah_h))
        draw_overlay.line([(0, y), (PAGE_W_PX, y)], fill=(0, 0, 0, alpha))
    overlay_bawah = overlay_bawah.filter(ImageFilter.GaussianBlur(radius=int(0.01 * PAGE_H_PX)))
    canvas.paste(overlay_bawah, (0, band_bawah_y0), overlay_bawah)
    draw = ImageDraw.Draw(canvas, 'RGBA')

    # ---- judul "PROFIL DESA" di tengah halaman ----
    f_eyebrow = _font(FONT_REGULAR, int(0.017 * PAGE_H_PX))
    f_judul_atas = _font(FONT_BOLD, int(0.044 * PAGE_H_PX))
    f_tahun = _font(FONT_REGULAR, int(0.019 * PAGE_H_PX))

    eyebrow = "P R O F I L   D E S A"
    judul_desa = desa.nama_desa.upper()
    teks_tahun = f"T A H U N   {desa.tahun_profil}"

    gap_eyebrow_judul = int(0.028 * PAGE_H_PX)
    gap_judul_tahun = int(0.040 * PAGE_H_PX)
    gap_tahun_garis = int(0.032 * PAGE_H_PX)
    tinggi_garis = max(2, int(0.004 * PAGE_H_PX))

    tinggi_eyebrow = f_eyebrow.getbbox(eyebrow)[3]
    tinggi_judul = f_judul_atas.getbbox(judul_desa)[3]
    tinggi_tahun = f_tahun.getbbox(teks_tahun)[3]

    tinggi_blok = (tinggi_eyebrow + gap_eyebrow_judul + tinggi_judul
                   + gap_judul_tahun + tinggi_tahun + gap_tahun_garis + tinggi_garis)

    # diposisikan di titik tengah keseluruhan halaman, bukan cuma area atas
    titik_tengah = int(PAGE_H_PX * 0.46)
    y_eyebrow = titik_tengah - tinggi_blok // 2

    bbox_eyebrow = draw.textbbox((0, 0), eyebrow, font=f_eyebrow)
    draw.text(((PAGE_W_PX - (bbox_eyebrow[2] - bbox_eyebrow[0])) / 2, y_eyebrow), eyebrow, font=f_eyebrow, fill='#0A3D62')

    y_judul = y_eyebrow + tinggi_eyebrow + gap_eyebrow_judul
    bbox_judul = draw.textbbox((0, 0), judul_desa, font=f_judul_atas)
    draw.text(((PAGE_W_PX - (bbox_judul[2] - bbox_judul[0])) / 2, y_judul), judul_desa, font=f_judul_atas, fill='#0A3D62')

    y_tahun = y_judul + tinggi_judul + gap_judul_tahun
    bbox_tahun = draw.textbbox((0, 0), teks_tahun, font=f_tahun)
    draw.text(((PAGE_W_PX - (bbox_tahun[2] - bbox_tahun[0])) / 2, y_tahun), teks_tahun, font=f_tahun, fill='#0A3D62')

    garis_y = y_tahun + tinggi_tahun + gap_tahun_garis
    garis_lebar = int(0.10 * PAGE_W_PX)
    draw.line(
        [(PAGE_W_PX / 2 - garis_lebar / 2, garis_y), (PAGE_W_PX / 2 + garis_lebar / 2, garis_y)],
        fill='#D9A253', width=tinggi_garis
    )

    # ---- "Terima Kasih..." di bagian bawah, polos tanpa background ----
    f_pesan = _font(FONT_BOLD, int(0.020 * PAGE_H_PX))
    t1 = "Terima Kasih atas Dukungan dan Partisipasi"
    t1b = "Seluruh Masyarakat"
    bbox1 = draw.textbbox((0, 0), t1, font=f_pesan)
    bbox1b = draw.textbbox((0, 0), t1b, font=f_pesan)

    jarak_antar_baris_pesan = int(0.030 * PAGE_H_PX)
    y_pesan_awal = band_bawah_y0 + int(0.055 * PAGE_H_PX)

    draw.text(((PAGE_W_PX - (bbox1[2] - bbox1[0])) / 2, y_pesan_awal), t1, font=f_pesan, fill='#F2F2F2')
    draw.text(((PAGE_W_PX - (bbox1b[2] - bbox1b[0])) / 2, y_pesan_awal + jarak_antar_baris_pesan), t1b, font=f_pesan, fill='#F2F2F2')

    # ---- logo + teks instansi, pojok kiri bawah ----
    logo_h = int(0.10 * PAGE_H_PX)   # <- UBAH DI SINI kalau logo cover belakang mau lebih besar/kecil
    logo_y = PAGE_H_PX - logo_h - int(0.032 * PAGE_H_PX)
    teks_x = pad

    if os.path.isfile(LOGO_PATH):
        logo = Image.open(LOGO_PATH).convert('RGBA')
        logo_w = int(logo.width * (logo_h / logo.height))
        logo_resized = logo.resize((logo_w, logo_h), Image.LANCZOS)
        canvas.paste(logo_resized, (pad, logo_y), logo_resized)
        teks_x = pad + logo_w + int(0.010 * PAGE_W_PX)

    lebar_maks_teks = PAGE_W_PX - teks_x - pad

    def _cari_ukuran_pas(teks, ukuran_awal, font_path):
        ukuran = ukuran_awal
        while ukuran > int(0.008 * PAGE_H_PX):
            font = _font(font_path, ukuran)
            lebar = draw.textbbox((0, 0), teks, font=font)[2]
            if lebar <= lebar_maks_teks:
                return font
            ukuran -= 1
        return _font(font_path, int(0.008 * PAGE_H_PX))

    f_pemda = _cari_ukuran_pas("PEMERINTAH KABUPATEN LAMPUNG TIMUR", int(0.019 * PAGE_H_PX), FONT_BOLD)
    f_desa = _cari_ukuran_pas(f"Pemerintah Desa {desa.nama_desa}", int(0.015 * PAGE_H_PX), FONT_REGULAR)

    tinggi_pemda = f_pemda.getbbox("PEMERINTAH")[3]
    tinggi_desa = f_desa.getbbox("Pemerintah")[3]
    jarak_antar_baris_logo = int(0.010 * PAGE_H_PX)
    tinggi_teks_total = tinggi_pemda + jarak_antar_baris_logo + tinggi_desa

    teks_y = logo_y + (logo_h - tinggi_teks_total) // 2

    draw.text((teks_x, teks_y), "PEMERINTAH KABUPATEN LAMPUNG TIMUR", font=f_pemda, fill='white')
    draw.text((teks_x, teks_y + tinggi_pemda + jarak_antar_baris_logo), f"Pemerintah Desa {desa.nama_desa}", font=f_desa, fill='#E7E7E7')

    out_path = os.path.join(GENERATED_DIR, f"cover_belakang_{desa.id}.png")
    canvas.save(out_path)
    return out_path


# ==================== DATA JAWABAN KUESIONER ====================

def get_peta_jawaban(desa):
    jawaban_qs = JawabanKuesioner.objects.filter(desa=desa).select_related('item')
    peta = {}
    for j in jawaban_qs:
        if j.nilai and j.nilai.strip():
            peta[j.item.no_item] = j.nilai.strip()
    return peta


# ==================== GAYA DASAR DOKUMEN (font, ukuran, spasi) ====================

def atur_gaya_dokumen(doc):
    _matikan_grid_dokumen(doc)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.line_spacing = 1.5   # <- line spacing paragraf isi (Normal)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    _matikan_snap_grid(normal)

    try:
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Times New Roman'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h1.paragraph_format.line_spacing = 1.5   # <- line spacing Heading 1 (judul bab)
        h1.paragraph_format.space_before = Pt(0)
        h1.paragraph_format.space_after = Pt(0)
        _matikan_snap_grid(h1)
    except KeyError:
        pass


    for nama_style_toc in ('TOC 1', 'TOC 2'):
        try:
            toc_style = doc.styles[nama_style_toc]
        except KeyError:
            toc_style = doc.styles.add_style(nama_style_toc, doc.styles['Normal'].type)
        toc_style.font.name = 'Times New Roman'
        toc_style.font.size = Pt(12)
        toc_style.font.color.rgb = RGBColor(0, 0, 0)
        toc_style.paragraph_format.line_spacing = 1.5
        toc_style.paragraph_format.space_before = Pt(0)
        toc_style.paragraph_format.space_after = Pt(4)

    try:
        h2 = doc.styles['Heading 2']
        h2.font.name = 'Times New Roman'
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2.paragraph_format.line_spacing = 1.5   # <- line spacing Heading 2 (judul tabel)
        h2.paragraph_format.space_before = Pt(6)
        h2.paragraph_format.space_after = Pt(3)
        _matikan_snap_grid(h2)
    except KeyError:
        pass


def _matikan_grid_dokumen(doc):
    """
    Nonaktifkan 'document grid' ala dokumen Asia Timur (China/Jepang) yang
    ikut terbawa dari template dasar python-docx. Grid ini penyebab jarak
    antar kata dan jarak antar paragraf/heading melebar tidak wajar.
    """
    for section in doc.sections:
        sectPr = section._sectPr
        docGrid = sectPr.find(qn('w:docGrid'))
        if docGrid is None:
            docGrid = OxmlElement('w:docGrid')
            sectPr.append(docGrid)
        docGrid.set(qn('w:type'), 'default')


def _matikan_snap_grid(style):
    """Pastikan paragraf pada style ini tidak lagi 'nempel' ke grid karakter/baris."""
    pPr = style.element.get_or_add_pPr()
    snap = pPr.find(qn('w:snapToGrid'))
    if snap is None:
        snap = OxmlElement('w:snapToGrid')
        pPr.append(snap)
    snap.set(qn('w:val'), '0')


def teks_hitam(paragraph, teks, bold=False, italic=False, size=12, align=None):
    run = paragraph.add_run(teks)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    if align:
        paragraph.alignment = align
    return run


def tambah_paragraf_body(doc, teks):
    """Paragraf isi: rata kanan-kiri, alinea menjorok, spasi 1,5."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    teks_hitam(p, teks, size=12)
    return p


def tambah_paragraf_body_multi(doc, teks):
    """
    Sama seperti tambah_paragraf_body, tapi untuk teks panjang hasil copy-paste
    dari PDF/Word (misal isian Sejarah Desa) yang rawan berantakan spasi/enter-nya.
    Teks dibersihkan dulu, lalu dipecah jadi beberapa paragraf Word yang rapi
    kalau memang ada baris kosong ganda (batas paragraf beneran) di teks aslinya.
    """
    teks_bersih = bersihkan_teks_tempel(teks)
    for paragraf in teks_bersih.split('\n\n'):
        if paragraf.strip():
            tambah_paragraf_body(doc, paragraf.strip())


def tambah_spacer(doc, jumlah=1):
    """Menambahkan paragraf kosong sebagai jarak/spasi antar elemen."""
    for _ in range(jumlah):
        doc.add_paragraph()


# ==================== FOOTER (nomor halaman + teks kiri) ====================

def tambah_footer(section, desa, label_dokumen="Profil Desa"):
    section.footer_distance = Cm(1.5)
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraf = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    paragraf.paragraph_format.tab_stops.clear_all()

    usable_width = section.page_width - section.left_margin - section.right_margin
    titik_tengah = usable_width // 2
    paragraf.paragraph_format.tab_stops.add_tab_stop(titik_tengah, WD_TAB_ALIGNMENT.CENTER)

    run_kiri = paragraf.add_run(f"{label_dokumen} {desa.nama_desa}")
    run_kiri.font.name = 'Times New Roman'
    run_kiri.font.size = Pt(10)
    run_kiri.italic = True
    run_kiri.font.color.rgb = RGBColor(0, 0, 0)

    paragraf.add_run("\t")

    run_page = paragraf.add_run()
    run_page.font.name = 'Times New Roman'
    run_page.font.size = Pt(10)
    run_page.font.color.rgb = RGBColor(0, 0, 0)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)


# ==================== HALAMAN GAMBAR PENUH (cover depan/belakang Profil Desa) ====================

def sisipkan_halaman_gambar_penuh(doc, image_path, is_first=False):
    if not is_first:
        doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.page_width = Cm(PAGE_W_CM)
    section.page_height = Cm(PAGE_H_CM)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(image_path, width=Cm(PAGE_W_CM), height=Cm(PAGE_H_CM))


def mulai_section_isi(doc, desa):
    """Section baru dengan margin seragam (normal untuk cetak/PDF) + footer untuk konten dokumen."""
    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.page_width = Cm(PAGE_W_CM)
    section.page_height = Cm(PAGE_H_CM)
    section.top_margin = Cm(3)      # <- margin halaman isi Profil Desa
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    tambah_footer(section, desa)
    return section


# ==================== HALAMAN PETA WILAYAH & KEPALA DESA (Profil Desa) ====================

def buat_halaman_peta(doc, desa):
    doc.add_heading(f"PETA WILAYAH ADMINISTRASI\n{desa.nama_desa.upper()}", level=1)
    tambah_spacer(doc, 1)
    if desa.peta_wilayah and os.path.isfile(desa.peta_wilayah.path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(desa.peta_wilayah.path, width=Cm(14))
    else:
        tambah_paragraf_body(doc, f"Peta wilayah administrasi {desa.nama_desa} belum tersedia dalam dokumen sumber.")
    doc.add_page_break()


def buat_halaman_kepala_desa(doc, desa):
    doc.add_heading(f"KEPALA DESA {desa.nama_desa.upper()}", level=1)
    tambah_spacer(doc, 1)

    if desa.nama_kepala_desa:
        p = doc.add_paragraph()
        teks_hitam(p, desa.nama_kepala_desa.upper(), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        tambah_spacer(doc, 1)

    if desa.foto_kepala_desa and os.path.isfile(desa.foto_kepala_desa.path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(desa.foto_kepala_desa.path, width=Cm(8))
    else:
        tambah_paragraf_body(doc, f"Foto Kepala Desa {desa.nama_desa} belum tersedia.")
    doc.add_page_break()


# ==================== KATA PENGANTAR (Profil Desa) ====================

def buat_kata_pengantar(doc, desa):
    doc.add_heading("KATA PENGANTAR", level=1)
    tambah_spacer(doc, 1)
    tambah_paragraf_body(doc,
        f"Puji dan syukur kami panjatkan kehadirat Tuhan Yang Maha Esa, atas rahmat dan karunia-Nya "
        f"penyusunan buku Profil Desa {desa.nama_desa} ini dapat terselesaikan."
    )
    tambah_paragraf_body(doc,
        f"Buku Profil Desa {desa.nama_desa} ini menggambarkan potensi dan gambaran umum perkembangan "
        f"desa yang disusun secara sistematis. Gambaran umum ini diperlukan sebagai referensi dan data "
        f"acuan dalam penyusunan program kegiatan pembangunan desa."
    )
    tambah_paragraf_body(doc,
        "Kami menyadari bahwa dalam menyediakan data dan informasi melalui buku ini masih banyak "
        "kekurangan. Oleh karena itu, kritik dan saran yang membangun sangat diharapkan untuk perbaikan "
        "buku ini. Harapan kami, semoga buku ini bermanfaat dan dapat dipergunakan sebagaimana mestinya."
    )
    tambah_spacer(doc, 1)
    p_ttd = doc.add_paragraph()
    teks_hitam(p_ttd, f"{desa.nama_desa}, {tanggal_indonesia()}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    p_jabatan = doc.add_paragraph()
    teks_hitam(p_jabatan, f"Kepala Desa {desa.nama_desa}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    tambah_spacer(doc, 3)
    if desa.nama_kepala_desa:
        p_nama = doc.add_paragraph()
        teks_hitam(p_nama, desa.nama_kepala_desa, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()


# ==================== DAFTAR ISI (dipakai Profil Desa & Monografi) ====================

def buat_daftar_isi(doc):
    doc.add_heading("DAFTAR ISI", level=1)
    tambah_spacer(doc, 1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = "Klik kanan lalu 'Update Field' untuk menampilkan daftar isi"
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2); run._r.append(placeholder); run._r.append(fld3)

    tambah_spacer(doc, 1)
    p_catatan = doc.add_paragraph()
    teks_hitam(p_catatan,
        "(Catatan: buka dokumen di Microsoft Word, tekan Ctrl+A lalu F9 untuk memperbarui daftar isi "
        "dan nomor halaman secara otomatis.)",
        italic=True, size=10
    )
    doc.add_page_break()


def buat_daftar_tabel(doc, semua_bab_data):
    doc.add_heading("DAFTAR TABEL", level=1)
    tambah_spacer(doc, 1)
    ada_tabel = False
    for bab, subbab_dengan_data in semua_bab_data:
        for subbab, _ in subbab_dengan_data:
            ada_tabel = True
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            teks_hitam(p, f"Tabel {subbab.kode_tabel}  {subbab.nama_tabel}", size=12)
    if not ada_tabel:
        tambah_paragraf_body(doc, "Tidak ada tabel yang ditampilkan pada dokumen ini.")
    doc.add_page_break()


# ==================== GAMBARAN UMUM (dipakai Profil Desa & Monografi) ====================

def buat_gambaran_umum(doc, desa, peta):
    doc.add_heading("GAMBARAN UMUM", level=1)
    tambah_spacer(doc, 1)
    status = peta.get(1)
    laut = peta.get(2)
    hutan = peta.get(8)

    kalimat = f"Secara administratif, {desa.nama_desa} "
    kalimat += f"berstatus sebagai {status.lower()}. " if status else "(data status wilayah tidak tersedia). "
    if laut:
        kalimat += "Desa ini berbatasan langsung dengan wilayah laut. " if laut.lower() == "ada" \
            else "Desa ini tidak berbatasan langsung dengan wilayah laut. "
    if hutan:
        kalimat += f"Lokasi desa berada {hutan.lower()} terhadap kawasan hutan."

    tambah_paragraf_body(doc, kalimat)
    if desa.sejarah_desa:
        tambah_paragraf_body_multi(doc, desa.sejarah_desa)
    doc.add_page_break()


# ==================== HEADER "DATA MONOGRAFI DESA DAN KELURAHAN" (khusus Monografi) ====================

def buat_data_monografi_header(doc, desa):
    """Header identitas 'DATA MONOGRAFI DESA DAN KELURAHAN', persis seperti dokumen asli --
    ditampilkan sebelum daftar Bidang A/B/C. Propinsi selalu 'Lampung' (hardcode)."""
    p_judul = doc.add_paragraph()
    teks_hitam(p_judul, "DATA MONOGRAFI DESA DAN KELURAHAN", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    tambah_spacer(doc, 1)

    baris = [
        ("1. Desa", desa.nama_desa),
        ("2. Nomor Kode", desa.nomor_kode or "-"),
        ("3. Kecamatan", desa.kecamatan),
        ("4. Kabupaten", desa.kabupaten),
        ("5. Propinsi", "Lampung"),   # <- hardcode, semua desa yang dikelola ada di Provinsi Lampung
        ("6. Keadaan Data", desa.keadaan_data or "-"),
    ]
    for label, nilai in baris:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.tab_stops.add_tab_stop(Cm(4.5))
        teks_hitam(p, f"{label}\t: {nilai}", size=12)
    tambah_spacer(doc, 1)


# ==================== TABEL DATA KUESIONER (abu-putih, header center) ====================

def _beri_warna_sel(cell, warna_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), warna_hex)
    tcPr.append(shd)


LEBAR_KOLOM_NO_CM = 1.2  # <- lebar kolom "No" pada tabel data kuesioner, ubah di sini kalau perlu


def _atur_lebar_dan_posisi_tabel(doc, table):
    """
    Buat lebar tabel fixed (bukan auto), rata tengah halaman, dan total lebar
    persis sama dengan area cetak (lebar halaman - margin kiri - margin kanan)
    dari section yang sedang aktif. Kolom "No" dibuat sempit & tetap,
    sisanya dibagi ke 2 kolom lain.
    """
    section = doc.sections[-1]
    lebar_area_cetak = section.page_width - section.left_margin - section.right_margin

    lebar_no = Cm(LEBAR_KOLOM_NO_CM)
    lebar_sisa = lebar_area_cetak - lebar_no
    lebar_kategori = int(lebar_sisa * 0.62)   # <- proporsi lebar kolom "Uraian" (dibesarin)
    lebar_nilai = lebar_sisa - lebar_kategori  # <- sisanya untuk kolom "Jumlah/Keterangan" (otomatis mengecil, ~38%)
    lebar_kolom = [lebar_no, lebar_kategori, lebar_nilai]

    table.autofit = False           # wajib False, kalau tidak Word abaikan lebar manual
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # tabel rata tengah halaman

    # Word membaca lebar per-sel di tiap baris, jadi harus di-set ulang di semua baris
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = lebar_kolom[idx]

    return lebar_kolom


def tambah_tabel(doc, baris_terisi, label_kategori="Uraian", label_nilai="Keterangan"):
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    # baris judul kolom: No | label_kategori | label_nilai
    hdr = table.rows[0].cells
    hdr[0].text = "No"
    hdr[1].text = label_kategori
    hdr[2].text = label_nilai

    # baris nomor urut kolom (1) (2) (3) -- ciri khas tabel statistik BPS
    sub = table.rows[1].cells
    sub[0].text = "(1)"
    sub[1].text = "(2)"
    sub[2].text = "(3)"

    for row_cells in (hdr, sub):
        for cell in row_cells:
            _beri_warna_sel(cell, 'D9D9D9')
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)

    for nomor, (nama_kolom, nilai) in enumerate(baris_terisi, start=1):
        row = table.add_row().cells
        row[0].text = str(nomor)
        row[1].text = nama_kolom
        row[2].text = nilai
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row:
            _beri_warna_sel(cell, 'FFFFFF')
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)

    _atur_lebar_dan_posisi_tabel(doc, table)


def kumpulkan_data_bab(bab, peta):
    subbab_dengan_data = []
    for subbab in bab.subbab_list.all().order_by('urutan'):
        baris_terisi = []
        for item_subbab in subbab.item_list.select_related('item').all():
            nilai = peta.get(item_subbab.item.no_item)
            if nilai:
                baris_terisi.append((item_subbab.nama_kolom, nilai))
        if baris_terisi:
            subbab_dengan_data.append((subbab, baris_terisi))
    return subbab_dengan_data


# ==================== HALAMAN CURRICULUM VITAE (BLANKO KOSONG, khusus Monografi) ====================

def _garis_bawah_paragraf(paragraph, ukuran=18, warna="000000"):
    """Tambah garis horizontal di bawah sebuah paragraf, dipakai sebagai penegas kop surat."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(ukuran))   # <- ketebalan garis, ubah di sini
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), warna)
    pBdr.append(bottom)
    pPr.append(pBdr)


def tambah_kop_surat_cv(doc, desa):
    """Kop surat resmi: logo di pojok kiri, teks Pemerintah Kabupaten -> Desa -> Kecamatan ->
    Sekretariat di sebelahnya (nempel bawah), diakhiri garis tunggal tebal.

    CATATAN PENTING soal garis penutup: paragraf p_garis di bawah SENGAJA dibuat
    "exactly 2pt" tingginya (bukan pakai ukuran font normal), karena kalau dibiarkan
    default, paragraf kosong tetap punya tinggi 1 baris penuh walau tanpa teks --
    itu yang bikin ada jarak "enter" aneh sebelum garis muncul. Jangan dihapus.
    """
    section = doc.sections[-1]
    lebar_area_cetak = section.page_width - section.left_margin - section.right_margin
    lebar_logo = int(lebar_area_cetak * 0.11)   # <- lebar kolom logo di kop surat (proporsi)
    lebar_teks = lebar_area_cetak - lebar_logo

    tabel_kop = doc.add_table(rows=1, cols=2)
    tabel_kop.autofit = False
    tabel_kop.allow_autofit = False
    tabel_kop.alignment = WD_TABLE_ALIGNMENT.CENTER

    sel_logo, sel_teks = tabel_kop.rows[0].cells
    sel_logo.width = lebar_logo
    sel_teks.width = lebar_teks

    # borderless: hilangkan semua garis tabel bawaan
    for cell in (sel_logo, sel_teks):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for sisi in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{sisi}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)

    sel_logo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p_logo = sel_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        p_logo.add_run().add_picture(LOGO_PATH, width=Cm(1.4))  

    sel_teks.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    p1 = sel_teks.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.line_spacing = 1.0
    teks_hitam(p1, f"PEMERINTAH KABUPATEN {desa.kabupaten.upper()}", bold=True, size=12)

    p2 = sel_teks.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.0
    teks_hitam(p2, f"DESA {desa.nama_desa.upper()}", bold=True, size=16)

    p3 = sel_teks.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.line_spacing = 1.0
    teks_hitam(p3, f"KECAMATAN {desa.kecamatan.upper()}", bold=True, size=12)

    p4 = sel_teks.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(0)
    p4.paragraph_format.space_after = Pt(0)
    p4.paragraph_format.line_spacing = 1.0
    teks_hitam(p4, "Sekretariat ..................................................", italic=True, size=10)

    # garis penutup kop surat -- tinggi paragraf dipaksa "exactly 2pt" biar nggak ada
    # jeda "1 baris kosong" sebelum garis (lihat catatan di docstring fungsi ini)
    p_garis = doc.add_paragraph()
    p_garis.paragraph_format.space_before = Pt(0)
    p_garis.paragraph_format.space_after = Pt(0)
    p_garis.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_garis.paragraph_format.line_spacing = Pt(2)
    _garis_bawah_paragraf(p_garis, ukuran=18)   # <- ketebalan garis kop surat, ubah di sini


def baris_isian(doc, nomor, label):
    """Satu baris identitas kosong dengan titik dua rata sejajar antar baris (pakai tab stop,
    bukan padding spasi, supaya benar-benar rata walau panjang label beda-beda)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.tab_stops.add_tab_stop(Cm(5.5))   # <- posisi titik dua, ubah di sini kalau label lebih panjang
    teks_hitam(p, f"{nomor}. {label}\t: " + "_" * 35, size=12)
    return p


def _atur_lebar_tabel_kustom(doc, table, proporsi):
    """Sama seperti _atur_lebar_dan_posisi_tabel, tapi menerima proporsi lebar per kolom
    sendiri (list pecahan yang totalnya 1) -- dipakai untuk tabel CV yang kolomnya beda-beda."""
    section = doc.sections[-1]
    lebar_area_cetak = section.page_width - section.left_margin - section.right_margin
    lebar_kolom = [int(lebar_area_cetak * p) for p in proporsi]

    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = lebar_kolom[idx]
    return lebar_kolom


def tambah_tabel_cv(doc, headers, proporsi, jumlah_baris_kosong=3):
    """Tabel isian kosong untuk halaman CV: header abu-abu tebal, baris data kosong bernomor
    (siap ditulis tangan atau diketik manual setelah dokumen didownload)."""
    kolom = len(headers)
    table = doc.add_table(rows=1, cols=kolom)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    for i, teks in enumerate(headers):
        hdr[i].text = teks
        _beri_warna_sel(hdr[i], 'D9D9D9')
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)

    for nomor in range(1, jumlah_baris_kosong + 1):   # <- jumlah baris kosong per tabel CV, ubah di pemanggilnya
        row = table.add_row().cells
        row[0].text = str(nomor)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in row:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)

    _atur_lebar_tabel_kustom(doc, table, proporsi)
    return table


def tambah_halaman_cv_kosong(doc, desa, judul_cv, label_jabatan, sertakan_keluarga=False):
    """
    Satu halaman blanko CV lengkap (kop surat + identitas + riwayat pendidikan/kursus/
    pekerjaan/organisasi + tanda tangan), semua masih kosong supaya bisa diisi manual
    oleh yang bersangkutan setelah dokumen di-download.

    sertakan_keluarga=True HANYA untuk CV Kepala Desa -- ini juga yang menentukan
    apakah blok tanda tangan di akhir halaman ikut ditampilkan atau tidak.
    """
    tambah_kop_surat_cv(doc, desa)
    tambah_spacer(doc, 1)

    p_judul = doc.add_paragraph()
    teks_hitam(p_judul, judul_cv, bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    tambah_spacer(doc, 1)

    # ---- I. IDENTITAS ----
    sub1 = doc.add_paragraph()
    teks_hitam(sub1, "I. IDENTITAS", bold=True, size=12)
    baris_isian(doc, 1, "Nama")
    baris_isian(doc, 2, "Tempat, Tanggal Lahir")
    baris_isian(doc, 3, "Kebangsaan")
    baris_isian(doc, 4, "Agama")
    baris_isian(doc, 5, "Pekerjaan")
    baris_isian(doc, 6, "Alamat")
    tambah_spacer(doc, 1)

    # ---- II. RIWAYAT PENDIDIKAN ----
    sub2 = doc.add_paragraph()
    teks_hitam(sub2, "II. RIWAYAT PENDIDIKAN", bold=True, size=12)
    tambah_tabel_cv(doc, ["No", "Pendidikan", "Tahun Lulus", "Keterangan"], [0.10, 0.45, 0.20, 0.25], 4)
    tambah_spacer(doc, 1)

    # ---- III. KURSUS/PELATIHAN ----
    sub3 = doc.add_paragraph()
    teks_hitam(sub3, "III. KURSUS/PELATIHAN YANG PERNAH DIIKUTI", bold=True, size=12)
    tambah_tabel_cv(doc, ["No", "Nama Kursus/Pelatihan", "Tahun", "Tempat", "Keterangan"], [0.08, 0.34, 0.14, 0.22, 0.22], 3)
    tambah_spacer(doc, 1)

    # ---- IV. RIWAYAT PEKERJAAN ----
    sub4 = doc.add_paragraph()
    teks_hitam(sub4, "IV. RIWAYAT PEKERJAAN", bold=True, size=12)
    tambah_tabel_cv(doc, ["No", "Pekerjaan", "Jabatan", "Tahun", "Keterangan"], [0.08, 0.32, 0.22, 0.16, 0.22], 3)
    tambah_spacer(doc, 1)

    # ---- V. ORGANISASI ----
    sub5 = doc.add_paragraph()
    teks_hitam(sub5, "V. ORGANISASI YANG PERNAH DIIKUTI", bold=True, size=12)
    tambah_tabel_cv(doc, ["No", "Nama Organisasi", "Jabatan", "Tahun", "Keterangan"], [0.08, 0.32, 0.22, 0.16, 0.22], 3)

    # ---- VI. DAFTAR KELUARGA + TANDA TANGAN (khusus Kepala Desa) ----
    if sertakan_keluarga:
        tambah_spacer(doc, 1)
        sub6 = doc.add_paragraph()
        teks_hitam(sub6, "VI. DAFTAR KELUARGA", bold=True, size=12)
        tambah_tabel_cv(doc, ["No", "Nama", "Tempat, Tanggal Lahir", "Jenis Kelamin", "Keterangan"], [0.07, 0.28, 0.28, 0.17, 0.20], 5)

        tambah_spacer(doc, 2)
        p_ttd1 = doc.add_paragraph()
        teks_hitam(p_ttd1, f"{desa.nama_desa}, ......................... 20....", align=WD_ALIGN_PARAGRAPH.RIGHT)
        p_ttd2 = doc.add_paragraph()
        teks_hitam(p_ttd2, f"{label_jabatan},", align=WD_ALIGN_PARAGRAPH.RIGHT)
        tambah_spacer(doc, 3)
        p_ttd3 = doc.add_paragraph()
        teks_hitam(p_ttd3, "(............................................)", align=WD_ALIGN_PARAGRAPH.RIGHT)


# ==================== GENERATOR UTAMA: PROFIL DESA ====================

def generate_profil_desa(desa):
    peta = get_peta_jawaban(desa)
    doc = Document()
    atur_gaya_dokumen(doc)

    section0 = doc.sections[0]
    section0.page_width = Cm(PAGE_W_CM)
    section0.page_height = Cm(PAGE_H_CM)

    cover_depan_path = generate_cover_depan(desa)
    sisipkan_halaman_gambar_penuh(doc, cover_depan_path, is_first=True)

    mulai_section_isi(doc, desa)

    buat_halaman_peta(doc, desa)
    buat_halaman_kepala_desa(doc, desa)
    buat_kata_pengantar(doc, desa)
    buat_daftar_isi(doc)

    semua_bab = list(Bab.objects.all().order_by('urutan'))
    semua_bab_data = [(bab, kumpulkan_data_bab(bab, peta)) for bab in semua_bab]

    buat_daftar_tabel(doc, semua_bab_data)
    buat_gambaran_umum(doc, desa, peta)

    # ---- isi Bab A s.d. K, tiap bab mulai halaman baru ----
    for i, (bab, subbab_dengan_data) in enumerate(semua_bab_data):
        if i > 0:
            doc.add_page_break()  # setiap bab baru wajib mulai halaman baru

        heading_bab = doc.add_heading(f"{bab.kode}. {bab.nama_bab.upper()}", level=1)
        heading_bab.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if not subbab_dengan_data:
            tambah_paragraf_body(doc, f"Tidak terdapat data mengenai {bab.nama_bab.lower()}.")
        else:
            for subbab, baris_terisi in subbab_dengan_data:
                judul_tabel = doc.add_heading(
                    f"Tabel {subbab.kode_tabel} {subbab.nama_tabel} di Desa {desa.nama_desa} Tahun {desa.tahun_profil}",
                    level=2
                )
                judul_tabel.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tambah_tabel(
                    doc, baris_terisi,
                    label_kategori=subbab.label_kategori,
                    label_nilai=subbab.label_nilai
                )

    cover_belakang_path = generate_cover_belakang(desa)
    sisipkan_halaman_gambar_penuh(doc, cover_belakang_path, is_first=False)

    return doc


# ==================== GENERATOR UTAMA: MONOGRAFI DESA ====================

def generate_monografi_desa(desa):
    peta = get_peta_jawaban(desa)
    doc = Document()
    atur_gaya_dokumen(doc)

    section0 = doc.sections[0]
    section0.page_width = Cm(PAGE_W_CM)
    section0.page_height = Cm(PAGE_H_CM)
    section0.top_margin = Cm(2)      # <- margin halaman Monografi
    section0.bottom_margin = Cm(2)
    section0.left_margin = Cm(2.5)
    section0.right_margin = Cm(2)
    tambah_footer(section0, desa, label_dokumen="Monografi Desa")

    # ---- COVER: MONOGRAFI / DESA-KECAMATAN-KABUPATEN / logo / TAHUN ----
    p_judul = doc.add_paragraph()
    p_judul.paragraph_format.space_before = Pt(40)
    teks_hitam(p_judul, "MONOGRAFI", bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER)

    p_sub1 = doc.add_paragraph()
    p_sub1.paragraph_format.space_before = Pt(30)
    teks_hitam(p_sub1, f"DESA {desa.nama_desa.upper()} KECAMATAN {desa.kecamatan.upper()}", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

    p_sub2 = doc.add_paragraph()
    teks_hitam(p_sub2, f"KABUPATEN {desa.kabupaten.upper()}", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

    p_logo = doc.add_paragraph()
    p_logo.paragraph_format.space_before = Pt(100)
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO_PATH):
        p_logo.add_run().add_picture(LOGO_PATH, width=Cm(6.5))   # <- UBAH DI SINI ukuran logo cover Monografi

    p_tahun = doc.add_paragraph()
    p_tahun.paragraph_format.space_before = Pt(100)
    teks_hitam(p_tahun, f"TAHUN {desa.tahun_profil}", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()
    buat_daftar_isi(doc)

    # ---- 4 halaman CV kosong ----
    tambah_halaman_cv_kosong(doc, desa, "CURRICULUM VITAE KEPALA DESA", "Kepala Desa", sertakan_keluarga=True)

    doc.add_page_break()
    tambah_halaman_cv_kosong(doc, desa, "CURRICULUM VITAE BPD", "Ketua BPD")

    doc.add_page_break()
    tambah_halaman_cv_kosong(doc, desa, "CURRICULUM VITAE LPMD", "Ketua LPMD")

    doc.add_page_break()
    tambah_halaman_cv_kosong(doc, desa, "CURRICULUM VITAE KETUA PKK", "Ketua PKK")

    # ---- Gambaran Umum + header Data Monografi ----
    doc.add_page_break()
    buat_gambaran_umum(doc, desa, peta)
    buat_data_monografi_header(doc, desa)

    # ---- isi Bidang A/B/C ----
    for bab in Bab.objects.all().order_by('urutan'):
        doc.add_heading(f"{bab.kode}. {bab.nama_bab.upper()}", level=2)
        subbab_dengan_data = kumpulkan_data_bab(bab, peta)
        if not subbab_dengan_data:
            tambah_paragraf_body(doc, f"Tidak terdapat data mengenai {bab.nama_bab.lower()}.")
        else:
            semua_baris = []
            for subbab, baris_terisi in subbab_dengan_data:
                semua_baris.extend(baris_terisi)
            tambah_tabel(doc, semua_baris)

        tambah_spacer(doc, 1)   # <- 1 enter jarak antar Bidang A/B/C, pakai spasi 1.5 dari style Normal (tanpa space before/after)

    return doc


# ==================== VIEW DOWNLOAD ====================

def _kirim_dokumen(doc, nama_file):
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{nama_file}"'
    doc.save(response)
    return response


def download_profil_desa(request, desa_id):
    desa = get_object_or_404(Desa, id=desa_id)
    doc = generate_profil_desa(desa)
    nama_file = f"Profil_Desa_{desa.nama_desa.replace(' ', '_')}_{desa.tahun_profil}.docx"
    return _kirim_dokumen(doc, nama_file)


def download_monografi_desa(request, desa_id):
    desa = get_object_or_404(Desa, id=desa_id)
    doc = generate_monografi_desa(desa)
    nama_file = f"Monografi_Desa_{desa.nama_desa.replace(' ', '_')}_{desa.tahun_profil}.docx"
    return _kirim_dokumen(doc, nama_file)